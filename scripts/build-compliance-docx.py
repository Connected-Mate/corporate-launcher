#!/usr/bin/env python3
"""Corporate Launcher — compliance .docx generator.

After ``scripts/generate.py`` has produced a launcher tree, this script
builds a polished Word document the creator can hand to their corporate
security office (RSSI / CISO / ANSSI correspondent). The document
summarizes the launcher's compliance posture across ~11 sections: cover
page, executive summary, architecture diagram, threat model, the 15
cyber controls, network perimeter, data classification, telemetry
posture, audit logging, offboarding plan, and a sign-off block.

Usage:
    python3 scripts/build-compliance-docx.py \
        --config config.json \
        --launcher-dir ~/.local/share/<slug> \
        [--audit-report report.json] \
        --out compliance.docx

External dependency:
    pip install python-docx

Pure Python 3.10+ otherwise (stdlib only).
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path
from typing import Any, Mapping

# --------------------------------------------------------------------- #
# Dependency detection                                                  #
# --------------------------------------------------------------------- #

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor
except ImportError:  # pragma: no cover — surfaced to the user
    sys.stderr.write(
        "error: python-docx is not installed.\n"
        "       install it with:  pip install python-docx\n"
    )
    sys.exit(1)


# --------------------------------------------------------------------- #
# Styling                                                               #
# --------------------------------------------------------------------- #

ACCENT = RGBColor(0x0B, 0x3D, 0x91)        # corporate dark blue
ACCENT_SOFT = RGBColor(0x4A, 0x6E, 0xB5)
GRAY_DARK = RGBColor(0x33, 0x33, 0x33)
GRAY_MID = RGBColor(0x66, 0x66, 0x66)
GRAY_LIGHT = RGBColor(0xD9, 0xD9, 0xD9)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# --------------------------------------------------------------------- #
# 15 cyber controls — extracted from templates/shared/cyber-rules.md.tpl #
# --------------------------------------------------------------------- #

CYBER_CONTROLS: list[tuple[str, str, str]] = [
    ("01", "HTTP security headers",
     "CSP without unsafe-inline/eval, HSTS, X-Content-Type-Options, Referrer-Policy, "
     "Permissions-Policy, X-Frame-Options DENY — browser-side defense in depth against XSS and clickjacking."),
    ("02", "Outdated components",
     "Refuse jQuery < 3.5, CKEditor 4, PHP < 8.1, any unmaintained lib. Pin latest stable. "
     "Specific versions carry known unpatched RCE/XSS CVEs."),
    ("03", "Cookie hardening",
     "HttpOnly + Secure + SameSite=Strict/Lax. Never SameSite=None. Session ≤ 8h active / 20min idle. "
     "Blocks document.cookie exfiltration and most CSRF."),
    ("04", "TLS posture",
     "Only TLS 1.2/1.3, ECDHE suites, SHA-256+, RSA ≥ 2048 or ECDSA P-256+. "
     "Deprecates BEAST/POODLE-era ciphers per RFC 8996."),
    ("05", "XSS prevention",
     "Escape every interpolated value. No innerHTML, outerHTML, eval, new Function, setTimeout(string). "
     "Each primitive parses a string as code or markup."),
    ("06", "SQL injection",
     "Parameterized queries only. Never string-concat user input into SQL. "
     "Prepared statements separate code from data at the protocol level."),
    ("07", "Authentication",
     "Lock account after 5 failed attempts. Generic error message ('invalid credentials'). "
     "Breaks online brute force and username enumeration."),
    ("08", "Session storage",
     "Server-side only. Never localStorage / sessionStorage for session tokens. "
     "Web Storage is reachable from any script on the origin."),
    ("09", "Password hashing",
     "bcrypt (cost ≥ 12) or Argon2id with unique salt. Never MD5 / SHA-1 / unsalted. "
     "Makes offline cracking economically infeasible after a leak."),
    ("10", "Logging discipline",
     "Zero secrets in logs. Zero stack traces in production responses. "
     "Log aggregators replicate to systems with broader access than the app."),
    ("11", "GDPR / CNIL",
     "Minimize collection, justify retention, support deletion. "
     "CNIL fines reach 4% of global revenue."),
    ("12", "Forbidden primitives",
     "No eval / new Function / innerHTML / setTimeout(string) / exec(userInput) / "
     "pickle.load / yaml.load / JSON-P. Each converts data into code."),
    ("13", "Input validation",
     "Always server-side. Client-side validation is UX only. "
     "The client is under attacker control (curl, intercepting proxy)."),
    ("14", "Secrets management",
     "Zero hardcoded credentials. Read from env vars or secret manager. "
     "Git history is forever; corporate scanners auto-revoke and page on-call."),
    ("15", "Forbidden patterns",
     "No SELECT * with user input, no SameSite=None, no unsafe-inline in CSP, "
     "no tokens in localStorage, no NODE_TLS_REJECT_UNAUTHORIZED=0 in prod."),
]


# --------------------------------------------------------------------- #
# Helpers                                                               #
# --------------------------------------------------------------------- #


def _set_cell_bg(cell, hex_rgb: str) -> None:
    """Set a table cell background colour (no python-docx native API)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tc_pr.append(shd)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             size: int | None = None, color: RGBColor | None = None,
             font: str = "Calibri") -> Any:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def _h1(doc, text: str) -> None:
    p = doc.add_heading(level=1)
    _add_run(p, text, bold=True, size=20, color=ACCENT, font="Calibri")


def _h2(doc, text: str) -> None:
    p = doc.add_heading(level=2)
    _add_run(p, text, bold=True, size=15, color=ACCENT_SOFT, font="Calibri")


def _h3(doc, text: str) -> None:
    p = doc.add_heading(level=3)
    _add_run(p, text, bold=True, size=12, color=GRAY_DARK, font="Calibri")


def _para(doc, text: str, *, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    _add_run(p, text, italic=italic, size=size, color=GRAY_DARK)


def _kv_row(table, key: str, value: str) -> None:
    row = table.add_row().cells
    row[0].text = ""
    row[1].text = ""
    _add_run(row[0].paragraphs[0], key, bold=True, size=10, color=GRAY_DARK)
    _add_run(row[1].paragraphs[0], value or "—", size=10, color=GRAY_DARK)
    _set_cell_bg(row[0], "F2F2F2")


def _yes_no(ctx: Mapping[str, Any], key: str) -> str:
    raw = str(ctx.get(key, "")).strip().lower()
    if raw in ("yes", "true", "1", "y"):
        return "Yes"
    if raw in ("no", "false", "0", "n", ""):
        return "No"
    return raw


# --------------------------------------------------------------------- #
# Sections                                                              #
# --------------------------------------------------------------------- #


def section_cover(doc, ctx: Mapping[str, Any]) -> None:
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, ctx.get("CORP_NAME", "Corporate Launcher"),
             bold=True, size=32, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, "Corporate Launcher — Compliance Dossier",
             size=16, color=GRAY_DARK)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(tagline, ctx.get("CORP_TAGLINE", ""), italic=True, size=12, color=GRAY_MID)

    for _ in range(3):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(meta, f"Prepared for: {ctx.get('CYBER_AUTHORITY', 'Corporate CISO')}",
             bold=True, size=12, color=GRAY_DARK)

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(org, f"Organization: {ctx.get('CORP_ORGANIZATION', '')}",
             size=11, color=GRAY_DARK)

    powered = doc.add_paragraph()
    powered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(powered, f"Powered by: {ctx.get('CORP_POWERED_BY', '')}",
             size=11, color=GRAY_MID)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(version, f"Launcher version: {ctx.get('CORP_LAUNCHER_VERSION', '1.0.0')}",
             size=11, color=GRAY_MID)

    when = doc.add_paragraph()
    when.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(when, f"Date: {date.today().isoformat()}",
             size=11, color=GRAY_MID)

    if ctx.get("RSSI_CLEARANCE_REF"):
        ref = doc.add_paragraph()
        ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(ref, f"Clearance reference: {ctx['RSSI_CLEARANCE_REF']}",
                 size=11, color=GRAY_MID)

    doc.add_page_break()


def section_toc(doc) -> None:
    _h1(doc, "Table of contents")
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    _para(doc, "(Right-click → Update Field in Word to refresh page numbers.)",
          italic=True, size=9)
    doc.add_page_break()


def section_executive_summary(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "1. Executive summary")
    clis = ", ".join(ctx.get("WRAPPED_CLIS", []) or []) or "(none)"
    backend = ctx.get("GATEWAY_BACKEND", ctx.get("CC_BACKEND", "corporate gateway"))
    summary = (
        f"{ctx.get('CORP_NAME', 'The launcher')} is a white-labeled wrapper around "
        f"the following coding assistants: {clis}. It routes every model call to "
        f"the corporate {backend} gateway at {ctx.get('CC_PRIMARY_URL', '(see config)')} "
        f"so that prompts, completions, and metadata stay inside the perimeter operated "
        f"by {ctx.get('CORP_ORGANIZATION', 'the organization')}. The launcher was "
        f"generated to give employees a sanctioned, auditable path to LLM-assisted "
        f"development while keeping the underlying vendor and any direct API endpoints "
        f"unreachable from user machines. This dossier maps the controls applied to "
        f"the threat model accepted by {ctx.get('CYBER_AUTHORITY', 'the corporate CISO')} "
        f"and supports the clearance request"
        f"{(' ' + ctx['RSSI_CLEARANCE_REF']) if ctx.get('RSSI_CLEARANCE_REF') else ''}."
    )
    _para(doc, summary)


def section_architecture(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "2. Architecture")
    _para(doc,
          "Data flow — each hop is enforced by configuration shipped with the "
          "launcher; the user has no path that bypasses the gateway.")
    diagram = [
        "  +-----------+        +---------------+        +------------------+        +-----------+",
        "  |   User    | -----> |   Launcher    | -----> |  Corp. Gateway   | -----> |    LLM    |",
        "  | (laptop)  |  TLS   | (white-label) |  mTLS  | (DLP + logging)  |  TLS   |  provider |",
        "  +-----------+        +---------------+        +------------------+        +-----------+",
        "                              |                        |",
        "                              v                        v",
        "                       Audit log (local)         SIEM / Splunk",
    ]
    p = doc.add_paragraph()
    _add_run(p, "\n".join(diagram), font="Consolas", size=9, color=GRAY_DARK)
    _para(doc,
          f"Gateway: {ctx.get('CC_PRIMARY_URL', '(configured)')} "
          f"({ctx.get('GATEWAY_BACKEND', 'corporate gateway')}). "
          f"Egress is constrained to this host via NO_PROXY={ctx.get('CORP_NO_PROXY', '')} "
          f"and the corporate proxy {ctx.get('PROXY_HOST', '')}:{ctx.get('PROXY_PORT', '')}.",
          italic=True, size=10)


def section_threat_model(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "3. Threat model")
    _para(doc, "The launcher addresses the following threats from the corporate risk register:")
    threats = [
        ("Data leak to vendor",
         "Prompts may contain customer names, IPs, draft code. Direct calls to public "
         "vendor APIs would exit the perimeter without DLP, DPA, or logging."),
        ("Telemetry exfiltration",
         "Upstream CLIs ship usage analytics, error reporters, voice modes, and "
         "feedback channels — each is a side channel that bypasses the gateway."),
        ("Vendor lock-in / identity drift",
         "Without the white-label identity, the assistant reveals the underlying model "
         "and provider, breaking the commercial agreement with the contracted gateway."),
        ("Unauthorized egress",
         "An auto-updater or marketplace fetch can pull arbitrary code from the public "
         "internet, bypassing the internal mirror and supply-chain controls."),
        ("Credential leak",
         "API tokens hardcoded or echoed back into Git, logs, or shell history are "
         "indistinguishable from a breach once they land in a replicated store."),
        ("Insecure code generation",
         "An assistant that emits eval, innerHTML, MD5 hashes, or unparameterized SQL "
         "becomes an automated source of CVEs in shipped products."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    _add_run(hdr[0].paragraphs[0], "Threat", bold=True, size=10, color=WHITE)
    _add_run(hdr[1].paragraphs[0], "Mitigation context", bold=True, size=10, color=WHITE)
    _set_cell_bg(hdr[0], "0B3D91")
    _set_cell_bg(hdr[1], "0B3D91")
    for name, why in threats:
        row = table.add_row().cells
        _add_run(row[0].paragraphs[0], name, bold=True, size=10, color=GRAY_DARK)
        _add_run(row[1].paragraphs[0], why, size=10, color=GRAY_DARK)


def section_cyber_controls(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "4. Cyber controls applied")
    _para(doc,
          f"These 15 controls are appended to every system prompt sent through "
          f"{ctx.get('CORP_NAME', 'the launcher')} via the rules file shipped at "
          f"{ctx.get('CORP_RULES_FILE', 'cyber-rules.md')}. They map to OWASP Top 10, "
          f"ANSSI secure-development guidance, and the corporate risk register.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, header in enumerate(("#", "Control", "Rationale")):
        _add_run(hdr[i].paragraphs[0], header, bold=True, size=10, color=WHITE)
        _set_cell_bg(hdr[i], "0B3D91")
    for num, name, rationale in CYBER_CONTROLS:
        row = table.add_row().cells
        _add_run(row[0].paragraphs[0], num, bold=True, size=9, color=ACCENT)
        _add_run(row[1].paragraphs[0], name, bold=True, size=10, color=GRAY_DARK)
        _add_run(row[2].paragraphs[0], rationale, size=9, color=GRAY_DARK)


def section_network_perimeter(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "5. Network perimeter")
    _para(doc, "Every outbound call is constrained by the settings below; the launcher "
               "refuses to start when the VPN probe fails or the CA bundle is missing.")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    _kv_row(table, "VPN required", _yes_no(ctx, "VPN_REQUIRED"))
    _kv_row(table, "VPN profile", ctx.get("VPN_PROFILE_NAME", ""))
    _kv_row(table, "VPN client", ctx.get("VPN_CLIENT_NAME", ""))
    _kv_row(table, "VPN probe URL", ctx.get("VPN_PROBE_URL", ""))
    _kv_row(table, "Proxy (HTTPS)", f"{ctx.get('PROXY_HOST', '')}:{ctx.get('PROXY_PORT', '')}")
    _kv_row(table, "NO_PROXY list", ctx.get("CORP_NO_PROXY", ""))
    _kv_row(table, "Corporate CA bundle", ctx.get("CA_BUNDLE_PATH", ""))
    _kv_row(table, "CA org issuer", ctx.get("CORP_CA_ORG", ""))
    _kv_row(table, "TLS inspection accepted", _yes_no(ctx, "ACCEPT_TLS_INSPECTION"))
    _kv_row(table, "Internal npm mirror", ctx.get("INTERNAL_NPM_MIRROR_URL", ""))


def section_data_classification(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "6. Data classification")
    _para(doc,
          f"Every prompt routed through {ctx.get('CORP_NAME', 'the launcher')} is treated "
          f"as Internal — Confidential by default. The launcher does not request, store, or "
          f"forward personal data; users are instructed via the system prompt to refuse pasting "
          f"secrets and to redirect to {ctx.get('CORP_SECRET_MANAGER', 'the corporate secret manager')}.")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    _kv_row(table, "Default classification", "Internal — Confidential")
    _kv_row(table, "PII expected", "No (out-of-scope by policy)")
    _kv_row(table, "Prompt filter active", _yes_no(ctx, "PROMPT_FILTER_ENABLED"))
    _kv_row(table, "Secret manager", ctx.get("CORP_SECRET_MANAGER", ""))
    _kv_row(table, "DPO contact", ctx.get("CORP_DPO_CONTACT", ""))
    _kv_row(table, "Applicable regimes", "GDPR, CNIL (FR), internal data charter")


def section_telemetry(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "7. Telemetry posture")
    _para(doc, "The launcher ships with every known upstream telemetry channel disabled. "
               "Each kill switch below is enforced via environment variables baked into "
               "the launcher wrapper, not via opt-in user setting.")
    switches = [
        ("Telemetry / analytics", _yes_no(ctx, "BLOCK_TELEMETRY"),
         "Disables vendor analytics, error reporting, usage pings."),
        ("Auto-update", _yes_no(ctx, "BLOCK_AUTO_UPDATE"),
         "Prevents the CLI from fetching arbitrary code from the public registry."),
        ("Feedback / bug-report commands", _yes_no(ctx, "BLOCK_FEEDBACK_CMDS"),
         "Disables /feedback, /report-bug and similar side channels."),
        ("Voice mode", _yes_no(ctx, "BLOCK_VOICE_MODE"),
         "Disables microphone capture and external speech APIs."),
        ("Forbidden vendor terms", "Enforced",
         f"System prompt blocks: {ctx.get('FORBIDDEN_TERMS', '')}"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Channel", "Disabled", "Notes")):
        _add_run(hdr[i].paragraphs[0], h, bold=True, size=10, color=WHITE)
        _set_cell_bg(hdr[i], "0B3D91")
    for name, state, notes in switches:
        row = table.add_row().cells
        _add_run(row[0].paragraphs[0], name, bold=True, size=10, color=GRAY_DARK)
        _add_run(row[1].paragraphs[0], state, size=10,
                 color=ACCENT if state in ("Yes", "Enforced") else GRAY_MID)
        _add_run(row[2].paragraphs[0], notes, size=9, color=GRAY_DARK)

    _para(doc, "Test evidence: run `scripts/smoke-test.sh` from the launcher tree; "
               "the script asserts that each blocked endpoint receives zero outbound "
               "connections during a 60-second canary session.", italic=True, size=10)


def section_audit_log(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "8. Audit log")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    _kv_row(table, "Audit system", ctx.get("CORP_AUDIT_SYSTEM", ""))
    _kv_row(table, "Sink location", ctx.get("CORP_AUDIT_LOCATION", ""))
    _kv_row(table, "Events captured",
            "session start/end, model + endpoint used, token counts, "
            "tool calls, refusal events, error class (no payload).")
    _kv_row(table, "Retention", "Per corporate SIEM policy (default 13 months)")
    _kv_row(table, "Access control",
            f"Read access restricted to {ctx.get('CYBER_AUTHORITY', 'CISO')} and "
            f"{ctx.get('CORP_INCIDENT_CONTACT', 'SOC')}; "
            "no employee can read logs of another employee.")
    _kv_row(table, "PII in logs", "None — payloads are not recorded, only metadata.")
    _kv_row(table, "Cost tracking",
            "Enabled" if _yes_no(ctx, "COST_TRACKING_ENABLED") == "Yes" else "Disabled")


def section_offboarding(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "9. Offboarding plan")
    _para(doc,
          f"Tokens issued to employees are bound to their {ctx.get('SSO_PROVIDER', 'SSO')} "
          f"identity and have a TTL of {ctx.get('TOKEN_TTL_DAYS', '30')} days. The full "
          f"revocation path is:")
    bullets = [
        f"HR triggers offboarding in {ctx.get('SSO_PROVIDER', 'SSO')}; the user's SSO "
        "identity is deactivated.",
        f"The gateway admin API ({ctx.get('GATEWAY_ADMIN_API', '')}) is called to revoke "
        "all tokens issued to that identity; subsequent launcher requests receive 401.",
        f"The token portal ({ctx.get('TOKEN_PORTAL_URL', '')}) stops issuing new tokens for "
        "the disabled identity automatically (SSO-gated).",
        "Local launcher state on the user's laptop becomes inert at the next start: the "
        "VPN probe still passes but every gateway call returns 401, and the launcher "
        "exits with a documented message.",
        f"For lost-laptop scenarios, {ctx.get('CORP_INCIDENT_CONTACT', 'SOC')} can revoke "
        "tokens out-of-band via the gateway admin API without waiting for HR.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        _add_run(p, b, size=11, color=GRAY_DARK)


def section_audit_findings(doc, audit: Mapping[str, Any] | None) -> None:
    if not audit:
        return
    _h1(doc, "10. Pre-submission audit findings")
    _para(doc, "The launcher tree was scanned before submission. Findings below "
               "are extracted verbatim from the audit report.")
    findings = audit.get("findings") or audit.get("issues") or []
    if not findings:
        _para(doc, "No findings — all checks passed.", italic=True)
        return
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Severity", "Check", "Detail")):
        _add_run(hdr[i].paragraphs[0], h, bold=True, size=10, color=WHITE)
        _set_cell_bg(hdr[i], "0B3D91")
    for f in findings:
        row = table.add_row().cells
        _add_run(row[0].paragraphs[0], str(f.get("severity", "info")),
                 bold=True, size=10, color=GRAY_DARK)
        _add_run(row[1].paragraphs[0], str(f.get("check", f.get("id", ""))),
                 size=10, color=GRAY_DARK)
        _add_run(row[2].paragraphs[0], str(f.get("detail", f.get("message", ""))),
                 size=9, color=GRAY_DARK)


def section_signoff(doc, ctx: Mapping[str, Any]) -> None:
    _h1(doc, "11. Sign-off")
    _para(doc,
          f"By signing below, {ctx.get('CYBER_AUTHORITY', 'the corporate CISO')} acknowledges "
          f"that {ctx.get('CORP_NAME', 'the launcher')} version "
          f"{ctx.get('CORP_LAUNCHER_VERSION', '1.0.0')} has been reviewed against the "
          f"controls described in this dossier and is cleared for internal distribution "
          f"under {ctx.get('CORP_LICENSE_NOTE', 'corporate licensing terms')}.")

    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    rows = table.rows
    labels = ("Name", "Role", "Date", "Signature")
    for i, label in enumerate(labels):
        cell_label, cell_value = rows[i].cells
        _add_run(cell_label.paragraphs[0], label, bold=True, size=11, color=GRAY_DARK)
        _set_cell_bg(cell_label, "F2F2F2")
        _add_run(cell_value.paragraphs[0], " ", size=11)

    doc.add_paragraph()
    _para(doc,
          f"For questions on this dossier: {ctx.get('CORP_INTERNAL_CONTACT', '')} | "
          f"Security incidents: {ctx.get('CORP_INCIDENT_CONTACT', '')} | "
          f"DPO: {ctx.get('CORP_DPO_CONTACT', '')}",
          italic=True, size=9)


# --------------------------------------------------------------------- #
# Document assembly                                                     #
# --------------------------------------------------------------------- #


def build_document(ctx: Mapping[str, Any], launcher_dir: Path,
                   audit: Mapping[str, Any] | None, out_path: Path) -> None:
    doc = Document()

    # Default style — corporate gray, Calibri base.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = GRAY_DARK

    # Page margins.
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # Footer with branding.
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(footer, f"{ctx.get('CORP_NAME', '')} — Compliance dossier — "
                         f"Powered by {ctx.get('CORP_POWERED_BY', '')}",
                 size=8, color=GRAY_MID)

    section_cover(doc, ctx)
    section_toc(doc)
    section_executive_summary(doc, ctx)
    section_architecture(doc, ctx)
    section_threat_model(doc, ctx)
    section_cyber_controls(doc, ctx)
    section_network_perimeter(doc, ctx)
    section_data_classification(doc, ctx)
    section_telemetry(doc, ctx)
    section_audit_log(doc, ctx)
    section_offboarding(doc, ctx)
    section_audit_findings(doc, audit)
    section_signoff(doc, ctx)

    # Appendix — pointer to the launcher tree on disk.
    _h2(doc, "Appendix — launcher tree")
    _para(doc, f"Generated launcher path: {launcher_dir}", size=10)
    _para(doc, f"Distribution mode: {ctx.get('DIST_MODE', 'none')}", size=10)
    _para(doc, f"Repository URL: {ctx.get('DIST_REPO_URL', '(local only)')}", size=10)
    _para(doc, f"Skills mode: {ctx.get('SKILLS_MODE', 'none')}", size=10)
    if ctx.get("MCP_SERVERS"):
        _para(doc, "MCP servers wired:", size=10)
        for srv in ctx["MCP_SERVERS"]:
            p = doc.add_paragraph(style="List Bullet")
            _add_run(p,
                     f"{srv.get('name')} — {srv.get('url')} "
                     f"(trusted: {'yes' if srv.get('trust') else 'no'})",
                     size=10, color=GRAY_DARK)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# --------------------------------------------------------------------- #
# CLI                                                                   #
# --------------------------------------------------------------------- #


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build the corporate compliance .docx for a generated launcher.",
    )
    parser.add_argument("--config", required=True, type=Path,
                        help="Path to the launcher config.json (same file passed to generate.py).")
    parser.add_argument("--launcher-dir", required=True, type=Path,
                        help="Path to the generated launcher tree (INSTALL_DIR).")
    parser.add_argument("--audit-report", type=Path, default=None,
                        help="Optional path to a JSON audit report to append.")
    parser.add_argument("--out", required=True, type=Path,
                        help="Output .docx path.")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)

    if not args.config.is_file():
        sys.stderr.write(f"error: config not found: {args.config}\n")
        return 2
    if not args.launcher_dir.exists():
        sys.stderr.write(f"warning: launcher directory does not exist: {args.launcher_dir}\n")

    try:
        ctx = json.loads(args.config.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        sys.stderr.write(f"error: invalid JSON in {args.config}: {exc}\n")
        return 2

    audit: Mapping[str, Any] | None = None
    if args.audit_report:
        if not args.audit_report.is_file():
            sys.stderr.write(f"warning: audit report not found: {args.audit_report}\n")
        else:
            try:
                audit = json.loads(args.audit_report.read_text(encoding="utf-8"))
            except json.JSONDecodeError as exc:
                sys.stderr.write(f"warning: invalid audit JSON ({exc}); skipping.\n")

    build_document(ctx, args.launcher_dir.expanduser(), audit, args.out)

    sys.stdout.write(f"wrote {args.out}\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
